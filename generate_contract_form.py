import abstra.workflows as aw
import os
import re
import shutil
import subprocess
from abstra.forms import *
from abstra.common import get_persistent_dir
from datetime import date
from unidecode import unidecode
from docxtpl import DocxTemplate
from docx import Document
from loguru import logger
import pypandoc

# getting data from workflow
register_dict = aw.get_data("register_info")

# In this code we are going to set a default document format with some tags to be filled
# the tags are defined by {{tag_name}}
# the document is going to be automatically filled with the user's data
# the user can also upload a model of a document and fill it with the tags

# regex syntax that will be used to fill the document
REGEX = r"\{\{(.*?)\}\}"

# formatts the label so that "_" are replaced by " "
# and the value so that it stays in unidecode mode
def transform_tags(tags):
    transformed = []
    for tag in tags:
        transformed.append(
            {
                "value": unidecode((tag.lower())),
                "label": tag.replace("_", " ").capitalize(),
            }
        )
    return transformed


def convert_with_pandoc(file_path):
    try:
        pypandoc.convert_file(file_path, "docx", outputfile=file_path)
        logger.success(f"Succesfully converted document: {file_path}")
        return None
    except OSError as e:
        logger.info(f"Error converting document with Pandoc: {e}")
        return None


def create_new_doc_with_tags(tags_values_dict, filepath, filename):
    doc = DocxTemplate(filepath)
    context = tags_values_dict
    try:
        doc.render(context)
    except Exception as e:
        problematic_tags = [k for k in context.keys() if k in str(e)]
        display(f"Error: {e}. Please check the following tags: {problematic_tags}")
    doc.save(filepath)
    return filepath


def generate_document(
    file_response, employee_national_id, contract_folder, contract_data={}
):
    # format the file name and the path
    filename = f"{date.today().strftime('%Y%m%d')}_{employee_national_id}"
    filepath = os.path.join(contract_folder, f"{filename}.docx")

    # creates a file on the path with the template on variable file_response
    with open(filepath, "wb") as out_file:
        out_file.write(file_response)

    # identify the parts of the archieve that matches with the regex syntax
    doc = Document(filepath)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    tags = re.findall(REGEX, text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_tags = re.findall(REGEX, cell.text)
                if cell_tags:
                    tags.append(cell_tags[0])
    tags_dict = dict.fromkeys(tags)
    tags_list = list(tags_dict.keys())
    tags_stripped_list = [t.strip() for t in tags_list]

    new_tags = [
        tag
        for tag in tags_stripped_list
        if tag not in list(contract_data.keys()) or contract_data[tag] is None
    ]

    new_tags_formatted = transform_tags(new_tags)
    tags_response = dict.fromkeys(tags_stripped_list)

    # for inputs in new_tags:
    new_tags_page = Page().display("Please fill the following data to generate the contract:")
    for tags in new_tags_formatted:
        label = tags["label"]
        new_tags_page = new_tags_page.read(f"{label}:", placeholder=label, key=tags["value"])

    if (len(new_tags_formatted) != 0):
        new_tags_page = new_tags_page.run()
        tags_response.update(dict(new_tags_page))
    
    # fill the contract with the data 
    tags_response.update(contract_data)
    output_filepath = create_new_doc_with_tags(tags_response, filepath, filename)

    # pandoc converts the file to a docx format
    convert_with_pandoc(output_filepath)

    display_file(
        output_filepath, download_text="Click here to download the filled document"
    )
    return output_filepath


def render(partial):
    if len(partial) != 0:
        if partial.get("existing_contract") == "New Contract":
            return Page().read_file("Upload your contract", key="contract_file", accepted_formats=[".docx"])


def render_upload(partial):
    if len(partial) != 0:
        if partial.get("approve") == "No":
            return Page().read_file("Upload the contract adjusted", key="adj_contract_file", accepted_formats=[".docx"])
        

template_overview_page = (
    Page()
    .display("Automatic Contract Generation", size="large")
    .display("For correct use, please upload a .docx template that follows this instructions", size="medium")
    .display_markdown('''
- On your template, use {{tag_name}} to define the fields that will be filled with the employee's data.
- Do not use spaces on the items to be replaced. Instead, use underscores (e.g. {{tag name}} needs to be written as {{tag_name}}).
- Do not use accents or special characters on the field's name.
- The {{tag_name}} must match the key in the dictionary that stores personal information in the workflow.
    ''')
    .run()
    )

contract_page = (
    Page()
    .read_dropdown(
        "Contract Model",
        ["New Contract", "Template Contract"],
        key="existing_contract"
    )
    .reactive(render)
    .run()
)

# gets the folder path to save the contracts
contract_folder = get_persistent_dir() / "contracts"

# if the folder doesnt exists, creates it 
contract_folder.mkdir(parents=True, exist_ok=True)

if register_dict["complement_address"] is None:
    register_dict["complement_address"] = ""

if (contract_page["existing_contract"] == "New Contract"):
    contract_file = contract_page["contract_file"].file.read()
    output_filepath = generate_document(contract_file, register_dict["identification_number"], contract_folder)
    aw.set_data("contract_data", {})
else:
    contract_file = open(f"contract_models/contract_template.docx", "rb").read()
    contract_data = {
        "started_at": register_dict["started_at"],
        "name": register_dict["name"],
        "position": register_dict["position"],
        "number_address": register_dict["number_address"],
        "complement_address": register_dict["complement_address"],
        "district": register_dict["district"],
        "address": register_dict["address"],
        "country": register_dict["country"],
        "zip_code": register_dict["zip_code"],
        "personal_email": register_dict["personal_email"],
        "internal_email": register_dict["internal_email"],
        "phone_number": register_dict["phone_number"],
        "salary": register_dict["salary"],
        "bank_name": register_dict["bank_name"],
        "bank_branch_code": register_dict["bank_branch_code"],
        "bank_account_number": register_dict["bank_account_number"],
        "birth_date": register_dict["birth_date"],
        "id_emitted_by": register_dict["id_emitted_by"],
        "identification_number": register_dict["identification_number"],
        "taxpayer_id": register_dict["taxpayer_id"]
    }

    aw.set_data("contract_data", contract_data)

    output_filepath = generate_document(contract_file, register_dict["identification_number"], contract_folder, contract_data)

approval_page = (
    Page()
    .display("Do you approve the previous contract?")
    .read_dropdown("Approve", ["Yes", "No"], key="approve")
    .reactive(render_upload)
    .run("Send")
)

if (approval_page["approve"] == "No"):
    open(output_filepath, "wb").write(approval_page["adj_contract_file"].file.read())

document_filename = "Service Agreement and Other Provisions"

aw.set_data(
    "contract_path",
    {
        "contract_filepath":output_filepath,
        "contract_filename":document_filename, 
    }
)
